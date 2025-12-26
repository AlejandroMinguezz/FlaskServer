import shutil
import cv2
import pytesseract
import os
from pdf2image import convert_from_path
import tempfile

def ejecutar_ocr(file_path: str, lang: str = "spa") -> str:
    try:
        # --- Localizar tesseract automáticamente ---
        tesseract_path = shutil.which("tesseract")

        # Si no está en PATH, buscar en rutas comunes de Windows
        if not tesseract_path:
            common_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                r"C:\Tesseract-OCR\tesseract.exe",
                "/usr/bin/tesseract",  # Linux/Docker
                "/usr/local/bin/tesseract"  # Linux/Docker
            ]
            for path in common_paths:
                if os.path.exists(path):
                    tesseract_path = path
                    print(f"[OCR] Tesseract encontrado en: {path}")
                    break

        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            # Configurar TESSDATA_PREFIX
            if "Windows" in os.name or os.name == "nt":
                base_dir = os.path.dirname(tesseract_path)
                os.environ["TESSDATA_PREFIX"] = os.path.join(base_dir, "tessdata")
        else:
            raise EnvironmentError("Tesseract no encontrado. Instala Tesseract OCR: https://github.com/UB-Mannheim/tesseract/wiki")

        # --- Procesamiento PDF ---
        if file_path.lower().endswith(".pdf"):
            paginas = convert_from_path(file_path, dpi=300)
            texto_total = []
            for pagina in paginas:
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                    pagina.save(tmp.name, "JPEG")
                    texto = ejecutar_ocr(tmp.name, lang)
                    texto_total.append(texto)
                    os.unlink(tmp.name)
            return " ".join(texto_total)

        # --- Procesamiento imagen ---
        img = cv2.imread(file_path)
        if img is None:
            raise FileNotFoundError(f"No se pudo leer la imagen: {file_path}")

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        gray = cv2.convertScaleAbs(gray, alpha=1.8, beta=10)
        gray = cv2.GaussianBlur(gray, (3, 3), 0)
        gray = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 5
        )

        config = "--psm 6 -c preserve_interword_spaces=1"
        text = pytesseract.image_to_string(gray, lang=lang, config=config)
        text = " ".join(text.split()).strip(" .,:;")

        print(f"[OCR] {os.path.basename(file_path)} → {len(text)} caracteres extraídos")
        return text

    except Exception as e:
        print(f"[ERROR OCR] {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """
    Extrae texto de archivos DOCX usando python-docx.
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extraer texto de párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        text = " ".join(text_parts)
        print(f"[DOCX] {os.path.basename(file_path)} → {len(text)} caracteres extraídos")
        return text

    except ImportError:
        print("[ERROR] python-docx no instalado. Ejecuta: pip install python-docx")
        return ""
    except Exception as e:
        print(f"[ERROR DOCX] {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """
    Extrae texto de archivos TXT planos.
    """
    try:
        # Intentar múltiples encodings comunes
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                print(f"[TXT] {os.path.basename(file_path)} → {len(text)} caracteres extraídos (encoding: {encoding})")
                return text
            except UnicodeDecodeError:
                continue

        # Si ningún encoding funcionó
        print(f"[ERROR TXT] No se pudo decodificar {file_path} con ningún encoding común")
        return ""

    except Exception as e:
        print(f"[ERROR TXT] {e}")
        return ""
