"""
DOCX text extraction module
"""

from docx import Document
from pathlib import Path


def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file

    Args:
        file_path: Path to the DOCX file

    Returns:
        dict: {
            'text': extracted text,
            'success': bool,
            'paragraphs': number of paragraphs,
            'error': error message if any
        }
    """
    try:
        file_path = Path(file_path)

        if not file_path.exists():
            return {
                'text': '',
                'success': False,
                'paragraphs': 0,
                'error': f'File not found: {file_path}'
            }

        doc = Document(file_path)

        # Extract text from paragraphs
        paragraphs_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                if row_text.strip():
                    tables_text.append(row_text)

        # Combine all text
        all_text = paragraphs_text + tables_text
        full_text = '\n'.join(all_text)

        return {
            'text': full_text,
            'success': True,
            'paragraphs': len(paragraphs_text),
            'tables': len(doc.tables),
            'error': None
        }

    except Exception as e:
        return {
            'text': '',
            'success': False,
            'paragraphs': 0,
            'error': str(e)
        }


def extract_text_from_docx_bytes(docx_bytes):
    """
    Extract text from DOCX file bytes

    Args:
        docx_bytes: DOCX file content as bytes

    Returns:
        dict: Same format as extract_text_from_docx
    """
    import io

    try:
        doc = Document(io.BytesIO(docx_bytes))

        # Extract text from paragraphs
        paragraphs_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                if row_text.strip():
                    tables_text.append(row_text)

        # Combine all text
        all_text = paragraphs_text + tables_text
        full_text = '\n'.join(all_text)

        return {
            'text': full_text,
            'success': True,
            'paragraphs': len(paragraphs_text),
            'tables': len(doc.tables),
            'error': None
        }

    except Exception as e:
        return {
            'text': '',
            'success': False,
            'paragraphs': 0,
            'error': str(e)
        }
